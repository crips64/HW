from __future__ import annotations

import base64
import json
import math
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OLD_REPORT = Path(r"C:\Users\annen\Desktop\Отчёт_5_Анненков_upd.docx")
NOTEBOOK = ROOT / "PSS_superpos_last_edt.ipynb"
OUT_DOCX = ROOT / "Отчёт_6_Анненков_PSS_superposition.docx"
ASSETS = ROOT / "report_assets"

SELECTED_MONTHS = [1, 5, 10, 120, 288]
ALL_MONTHS = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 18, 24, 30, 36, 42, 48, 54, 60, 66, 72, 78, 84, 96, 108, 120, 288]

COMPARISON_ROWS = [
    (1, 217.843044, 217.822756, 0.020288, 0.081318, 25.358385, 0.645567),
    (2, 217.840822, 217.803901, 0.036921, 0.093465, 24.797841, 0.623034),
    (3, 217.838599, 217.783073, 0.055526, 0.108847, 24.424766, 0.617032),
    (4, 217.836377, 217.762925, 0.073452, 0.124433, 24.171075, 0.619962),
    (5, 217.834155, 217.742095, 0.092060, 0.141081, 23.971294, 0.627904),
    (6, 217.831933, 217.721918, 0.110016, 0.157402, 23.813643, 0.638341),
    (7, 217.829712, 217.701041, 0.128671, 0.174556, 23.676224, 0.650918),
    (8, 217.827490, 217.680133, 0.147357, 0.191873, 23.556566, 0.664671),
    (9, 217.825269, 217.659865, 0.165403, 0.208687, 23.453079, 0.678723),
    (10, 217.823047, 217.638884, 0.184164, 0.226252, 23.356315, 0.693785),
    (11, 217.820826, 217.618540, 0.202286, 0.243277, 23.270004, 0.708695),
    (12, 217.818605, 217.597474, 0.221131, 0.261044, 23.187223, 0.724351),
    (18, 217.805317, 217.473515, 0.331802, 0.366398, 22.781496, 0.816056),
    (24, 217.792131, 217.345630, 0.446501, 0.476633, 22.438777, 0.906412),
    (30, 217.779084, 217.217840, 0.561243, 0.588378, 22.131994, 0.993760),
    (36, 217.766186, 217.085824, 0.680362, 0.705537, 21.842379, 1.082289),
    (42, 217.753445, 216.953105, 0.800340, 0.824142, 21.588177, 1.171689),
    (48, 217.740863, 216.816738, 0.924126, 0.946233, 21.390444, 1.269411),
    (54, 217.728439, 216.680434, 1.048006, 1.068792, 21.289955, 1.374572),
    (60, 217.716161, 216.539666, 1.176495, 1.196669, 21.302911, 1.490038),
    (66, 217.704015, 216.399106, 1.304909, 1.325192, 21.441737, 1.609171),
    (72, 217.691984, 216.254243, 1.437742, 1.457369, 21.677430, 1.733922),
    (78, 217.680057, 216.109797, 1.570260, 1.588462, 21.981336, 1.859334),
    (84, 217.668220, 215.960978, 1.707242, 1.722677, 22.303003, 1.989517),
    (96, 217.644845, 215.659050, 1.985795, 1.994391, 22.882962, 2.255896),
    (108, 217.621701, 215.350170, 2.271531, 2.273378, 23.264323, 2.530473),
    (120, 217.598737, 215.034233, 2.564504, 2.560594, 23.424132, 2.812574),
    (288, 217.292206, 210.091269, 7.200936, 7.089595, 17.185844, 7.319235),
]


def font(size: int = 22, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def read_grid(path: Path, skiprows: int = 6) -> list[list[float]]:
    rows = []
    with path.open("r", encoding="utf-8", errors="ignore") as fh:
        for line in fh:
            if line.startswith("#") or not line.strip():
                continue
            parts = line.split()
            if len(parts) < 5:
                continue
            value = float(parts[2])
            col = int(float(parts[3])) - 1
            row = int(float(parts[4])) - 1
            rows.append((row, col, value))
    nrow = max(r for r, _, _ in rows) + 1
    ncol = max(c for _, c, _ in rows) + 1
    grid = [[0.0 for _ in range(ncol)] for _ in range(nrow)]
    for r, c, v in rows:
        grid[r][c] = v
    return grid


def viridis_like(t: float) -> tuple[int, int, int]:
    t = min(1.0, max(0.0, t))
    stops = [
        (0.00, (68, 1, 84)),
        (0.25, (59, 82, 139)),
        (0.50, (33, 145, 140)),
        (0.75, (94, 201, 98)),
        (1.00, (253, 231, 37)),
    ]
    for idx in range(len(stops) - 1):
        a_t, a_c = stops[idx]
        b_t, b_c = stops[idx + 1]
        if a_t <= t <= b_t:
            local = (t - a_t) / (b_t - a_t)
            return tuple(int(a_c[k] + local * (b_c[k] - a_c[k])) for k in range(3))
    return stops[-1][1]


def draw_heatmap_panel(grid: list[list[float]], title: str, subtitle: str = "", log_scale: bool = False) -> Image.Image:
    values = [v for row in grid for v in row if math.isfinite(v)]
    if log_scale:
        transformed = [[math.log10(max(v, 1e-6)) for v in row] for row in grid]
        values = [v for row in transformed for v in row]
    else:
        transformed = grid
    vmin, vmax = min(values), max(values)
    h, w = len(grid), len(grid[0])
    scale = max(1, 320 // max(w, h))
    img = Image.new("RGB", (w, h), "white")
    pix = img.load()
    for r in range(h):
        for c in range(w):
            t = 0.0 if vmax == vmin else (transformed[r][c] - vmin) / (vmax - vmin)
            pix[c, h - 1 - r] = viridis_like(t)
    img = img.resize((w * scale, h * scale), Image.Resampling.NEAREST)
    panel = Image.new("RGB", (img.width + 90, img.height + 82), "white")
    d = ImageDraw.Draw(panel)
    d.text((0, 0), title, fill=(30, 30, 30), font=font(18, bold=True))
    if subtitle:
        d.text((0, 24), subtitle, fill=(80, 80, 80), font=font(13))
    panel.paste(img, (0, 50))
    bar_x = img.width + 18
    bar_y = 50
    for y in range(img.height):
        t = 1 - y / max(1, img.height - 1)
        d.line((bar_x, bar_y + y, bar_x + 14, bar_y + y), fill=viridis_like(t))
    d.rectangle((bar_x, bar_y, bar_x + 14, bar_y + img.height), outline=(60, 60, 60))
    shown_min, shown_max = (10**vmin, 10**vmax) if log_scale else (vmin, vmax)
    d.text((bar_x + 20, bar_y), f"{shown_max:.2g}", fill=(40, 40, 40), font=font(11))
    d.text((bar_x + 20, bar_y + img.height - 14), f"{shown_min:.2g}", fill=(40, 40, 40), font=font(11))
    return panel


def make_model_maps() -> Path:
    panels = [
        draw_heatmap_panel(read_grid(ROOT / "data" / "case_3" / "poro_integral_3"), "Пористость", "доли ед."),
        draw_heatmap_panel(read_grid(ROOT / "data" / "case_3" / "permx_integral_3"), "Проницаемость", "мД, лог. шкала", True),
        draw_heatmap_panel(read_grid(ROOT / "data" / "heff"), "Эффективная толщина", "м"),
    ]
    gutter = 26
    out = Image.new("RGB", (sum(p.width for p in panels) + gutter * 2, max(p.height for p in panels)), "white")
    x = 0
    for p in panels:
        out.paste(p, (x, 0))
        x += p.width + gutter
    path = ASSETS / "model_maps.png"
    out.save(path)
    return path


def extract_notebook_figures() -> dict[int, Path]:
    existing_paths = {month: ASSETS / f"comparison_{month:03d}.png" for month in SELECTED_MONTHS}
    if all(path.exists() for path in existing_paths.values()):
        return existing_paths

    nb = json.loads(NOTEBOOK.read_text(encoding="utf-8"))
    image_outputs = [o for o in nb["cells"][89].get("outputs", []) if "image/png" in (o.get("data") or {})]
    if len(image_outputs) != len(ALL_MONTHS):
        raise RuntimeError(f"Expected {len(ALL_MONTHS)} comparison figures, got {len(image_outputs)}")

    paths = {}
    for month in SELECTED_MONTHS:
        idx = ALL_MONTHS.index(month)
        data = image_outputs[idx]["data"]["image/png"]
        if isinstance(data, list):
            data = "".join(data)
        path = ASSETS / f"comparison_{month:03d}.png"
        path.write_bytes(base64.b64decode(data))
        paths[month] = path
    return paths


def draw_error_chart() -> Path:
    width, height = 1100, 620
    margin_l, margin_t, margin_r, margin_b = 95, 82, 35, 95
    plot_w = width - margin_l - margin_r
    plot_h = height - margin_t - margin_b
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    title_font = font(28, True)
    axis_font = font(18)
    small_font = font(15)
    d.text((margin_l, 16), "Динамика ошибки PSS-DTOF относительно tNavigator", fill=(25, 25, 25), font=title_font)

    months = [row[0] for row in COMPARISON_ROWS]
    mean_abs = [row[4] for row in COMPARISON_ROWS]
    rmse = [row[6] for row in COMPARISON_ROWS]
    max_y = max(max(mean_abs), max(rmse)) * 1.08
    min_x, max_x = min(months), max(months)

    def xp(m: float) -> float:
        return margin_l + (m - min_x) / (max_x - min_x) * plot_w

    def yp(v: float) -> float:
        return margin_t + plot_h - v / max_y * plot_h

    d.rectangle((margin_l, margin_t, margin_l + plot_w, margin_t + plot_h), outline=(80, 80, 80))
    for tick in [0, 1, 2, 3, 4, 5, 6, 7, 8]:
        y = yp(tick)
        d.line((margin_l, y, margin_l + plot_w, y), fill=(225, 225, 225))
        d.text((40, y - 9), str(tick), fill=(55, 55, 55), font=small_font)
    for tick in [1, 12, 60, 120, 180, 240, 288]:
        x = xp(tick)
        d.line((x, margin_t + plot_h, x, margin_t + plot_h + 6), fill=(60, 60, 60))
        d.text((x - 16, margin_t + plot_h + 12), str(tick), fill=(55, 55, 55), font=small_font)

    def draw_line(vals: list[float], color: tuple[int, int, int]) -> None:
        points = [(xp(m), yp(v)) for m, v in zip(months, vals)]
        d.line(points, fill=color, width=4)
        for x, y in points:
            d.ellipse((x - 4, y - 4, x + 4, y + 4), fill=color)

    draw_line(mean_abs, (44, 127, 184))
    draw_line(rmse, (214, 92, 59))
    d.text((margin_l + 10, margin_t + 16), "средняя |Δp|", fill=(44, 127, 184), font=axis_font)
    d.text((margin_l + 10, margin_t + 42), "RMSE", fill=(214, 92, 59), font=axis_font)
    d.text((margin_l + plot_w // 2 - 60, height - 38), "Время, мес.", fill=(40, 40, 40), font=axis_font)
    d.text((margin_l, margin_t - 26), "Ошибка, бар", fill=(40, 40, 40), font=axis_font)
    path = ASSETS / "error_dynamics.png"
    img.save(path)
    return path


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)


def set_style_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 15), ("Heading 3", 14)]:
        if style_name in doc.styles:
            st = doc.styles[style_name]
            st.font.name = "Times New Roman"
            st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            st.font.size = Pt(size)
            st.font.bold = True
            st.font.color.rgb = RGBColor(0, 0, 0)
            st.paragraph_format.first_line_indent = Cm(0)
            st.paragraph_format.space_before = Pt(12)
            st.paragraph_format.space_after = Pt(6)


def add_centered(doc: Document, text: str, size: int = 14, bold: bool = False, spacing_after: int = 0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(spacing_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    return p


def add_body_paragraph(doc: Document, text: str = "", bold_lead: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.15
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(14)
        text = text[len(bold_lead):]
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(14)
    return p


def add_formula(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    o_math_para = OxmlElement("m:oMathPara")
    o_math_para_pr = OxmlElement("m:oMathParaPr")
    jc = OxmlElement("m:jc")
    jc.set(qn("m:val"), "center")
    o_math_para_pr.append(jc)
    o_math_para.append(o_math_para_pr)

    o_math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    run_props = OxmlElement("m:rPr")
    math_run.append(run_props)
    math_text = OxmlElement("m:t")
    math_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    math_text.text = text
    math_run.append(math_text)
    o_math.append(math_run)
    o_math_para.append(o_math)
    p._p.append(o_math_para)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        r = p.add_run("- ")
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(14)
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(14)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)


def add_picture(doc: Document, path: Path, width_cm: float, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_comparison_table(doc: Document) -> None:
    selected = [row for row in COMPARISON_ROWS if row[0] in SELECTED_MONTHS]
    headers = ["Месяц", "PSS, бар", "tNavigator, бар", "Δpср, бар", "mean |Δp|, бар", "max |Δp|, бар", "RMSE, бар"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], "D9EAF7")
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=9)
    for row in selected:
        cells = table.add_row().cells
        values = [str(row[0])] + [f"{x:.3f}" for x in row[1:]]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value, size=9)
    add_caption(doc, "Таблица 1 - Метрики ошибки PSS-DTOF относительно tNavigator для выбранных сроков")


def add_toc(doc: Document) -> None:
    add_heading(doc, "Содержание", level=1)
    items = [
        "Реферат",
        "Введение",
        "1. Модель карбонатного коллектора Costa",
        "2. Внесенные изменения в реализацию суперпозиции",
        "3. Результаты вычислений и сравнение с tNavigator",
        "4. Выводы",
        "Список используемых источников",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(14)


def build_doc(model_maps: Path, comparison_paths: dict[int, Path], error_chart: Path) -> None:
    doc = Document(str(OLD_REPORT))
    clear_document(doc)
    set_margins(doc)
    set_style_defaults(doc)

    add_centered(doc, "Министерство науки и высшего образования Российской Федерации", 12)
    add_centered(doc, "федеральное государственное автономное образовательное учреждение", 12)
    add_centered(doc, "высшего образования", 12)
    add_centered(doc, "НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ", 12, True)
    add_centered(doc, "ТОМСКИЙ ПОЛИТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ", 12, True)
    add_centered(doc, "(ФГАОУ ВО НИ ТПУ)", 12)
    for _ in range(4):
        doc.add_paragraph()
    add_centered(doc, "ОТЧЕТ", 16, True)
    add_centered(doc, "О НАУЧНО-ИССЛЕДОВАТЕЛЬСКОЙ РАБОТЕ", 14, True, 12)
    add_centered(
        doc,
        "Применение PSS-DTOF суперпозиции для 12-скважинной модели карбонатного коллектора Costa и сравнение с tNavigator",
        14,
        True,
        10,
    )
    add_centered(doc, "по теме:", 12)
    add_centered(doc, "«Цифровая нефтесервисная компания: цифровое сопровождение добычи и переработки нефти и газа»", 12)
    for _ in range(5):
        doc.add_paragraph()
    add_body_paragraph(doc, "Руководитель проекта __________________ В.С. Рукавишников")
    add_body_paragraph(doc, "Исполнитель: Анненков Илья Сергеевич")
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "Томск 2026", 12)
    doc.add_page_break()

    add_heading(doc, "Список исполнителей", level=1)
    add_body_paragraph(doc, "Исполнитель: Анненков Илья Сергеевич")
    doc.add_page_break()

    add_heading(doc, "Реферат", level=1)
    add_body_paragraph(doc, "В данной работе продолжается развитие подхода псевдоустановившегося давления (PSS) и диффузионного времени пролета (DTOF) для ускоренного гидродинамического моделирования карбонатного коллектора Costa.")
    add_body_paragraph(doc, "Рассмотрена постановка суперпозиции для группы из 12 скважин. Расчет давления выполнен по PSS-DTOF координате и сопоставлен с результатами полноразмерной гидродинамической симуляции в tNavigator.")
    add_body_paragraph(doc, "Ключевые слова: гидродинамическое моделирование, PSS, DTOF, суперпозиция, карбонатный коллектор, Costa, tNavigator, давление, ошибка моделирования.")
    doc.add_page_break()

    add_toc(doc)
    doc.add_page_break()

    add_heading(doc, "Введение", level=1)
    add_body_paragraph(doc, "Метод псевдоустановившегося давления (Pseudo-Steady State, PSS) используется для построения пространственной координаты, связанной с распределением давления при единичном дебите. Подход DTOF (Diffusive Time of Flight) позволяет учитывать не только пространственное положение ячеек, но и характер вовлечения порового объема во времени.")
    add_body_paragraph(doc, "В предыдущей работе основное внимание уделялось расчету для одной скважины и псевдоустановившемуся режиму. В текущей работе рассматривается расширение подхода на группу из 12 скважин с использованием суперпозиции вкладов давления от каждой скважины.")
    add_body_paragraph(doc, "Цель работы - реализовать расчет PSS-DTOF суперпозиции для 12-скважинной модели карбонатного коллектора Costa и сопоставить полученные карты давления с результатами tNavigator.")
    add_body_paragraph(doc, "Для достижения цели в работе решаются следующие задачи:")
    add_bullets(doc, [
        "описать модель коллектора Costa и исходные карты фильтрационно-емкостных свойств;",
        "перейти от односкважинной постановки к матрице взаимного влияния 12 скважин;",
        "включить в расчет формулы, учитывающие начальный и переходный режимы фильтрации;",
        "сравнить карты давления PSS-DTOF и tNavigator для выбранных моментов времени;",
        "оценить среднюю, максимальную и RMSE-ошибку давления.",
    ])

    add_heading(doc, "1. Модель карбонатного коллектора Costa", level=1)
    add_body_paragraph(doc, "В качестве объекта исследования используется карбонатный коллектор Costa. Модель характеризуется выраженной неоднородностью фильтрационных свойств и наличием трещин, что делает ее удобной для проверки ускоренных методов моделирования.")
    add_body_paragraph(doc, "В расчетной постановке используются осредненные по толщине карты пористости, проницаемости и эффективной толщины. Эти карты задают распределение порового объема и проводимости, которые далее применяются при построении PSS- и DTOF-координат.")
    add_picture(doc, model_maps, 16.5, "Рисунок 1 - Карты пористости, проницаемости и эффективной толщины модели Costa")
    add_body_paragraph(doc, "По классификации Nelson карбонатные трещинные резервуары разделяются по роли матрицы и трещин в накоплении и фильтрации. В одном предельном случае матрица обеспечивает основной объем хранения, а трещины являются каналами повышенной проводимости. В другом случае вклад трещин становится определяющим как для проводимости, так и для емкости. Модель Costa относится к классу объектов, где учет трещинной составляющей важен для корректного воспроизведения фильтрационного отклика.")
    add_body_paragraph(doc, "В этом разделе можно дополнительно раскрыть геологическое описание Costa, привести параметры сетки и добавить более подробную связь с классификацией Nelson.")

    add_heading(doc, "2. Внесенные изменения в реализацию суперпозиции", level=1)
    add_body_paragraph(doc, "В предыдущем отчете суперпозиция была записана в общем интегральном виде для двух скважин. Там уже были введены отдельные DTOF-системы координат для каждой скважины, сумма вкладов давления, коэффициенты M_ij и переключение режимов скважин через индикатор F_i. Эти формулы в текущем отчете повторно не приводятся, поскольку их физический смысл не изменился.")
    add_body_paragraph(doc, "Основное изменение текущей реализации состоит не в замене принципа суперпозиции, а в его вычислительной форме. Алгоритм был перенесен с двухскважинной демонстрационной схемы на 12 скважин модели Costa. Для каждой скважины строится собственная карта τ_j(x), а точка влияния другой скважины определяется как номер интервала на этой карте.")
    add_body_paragraph(doc, "Первое содержательное отличие связано с учетом переходного режима. В прошлой работе основной расчет был ориентирован на псевдоустановившийся отклик, а сейчас безразмерный поток считается как функция времени и вовлеченного порового объема:")
    add_formula(doc, "Wₖ(t) = Σₘ≥ₖ exp(-τₘ² / 4t) ΔVₚ,ₘ")
    add_formula(doc, "qD,ₖ(t) = Wₖ(t) / W₀(t)")
    add_body_paragraph(doc, "Далее интегральный вклад из старой записи M_ij заменен накопленной дискретной функцией A_k(t). Она считается отдельно для каждой скважины-источника и каждого временного шага:")
    add_formula(doc, "Aₖ(t) = Aₖ₋₁(t) + μ [ qD,ₖ₋₁(t)/Tₖ⁻ + qD,ₖ(t)/Tₖ⁺ ]")
    add_body_paragraph(doc, "После этого коэффициент взаимного влияния между скважиной-наблюдателем i и скважиной-источником j записывается в короткой дискретной форме:")
    add_formula(doc, "Mᵢⱼ(t) = Aⱼ(kᵢⱼ,t) - Jⱼ⁻¹(t)")
    add_body_paragraph(doc, "Здесь k_ij - номер интервала на τ-карте скважины j, в который попадает скважина i. Величина J_j^{-1}(t) рассчитывается в коде как среднее значение накопленного сопротивления по вовлеченному поровому объему; отдельную интегральную формулу из прошлого отчета здесь не дублируем.")
    add_body_paragraph(doc, "Вторая практическая доработка связана с решением системы. Вместо явной блочной матрицы с индикаторами F_i в коде используется список controls: для каждой скважины задается режим rate или bhp. Если при заданном дебите расчетное забойное давление становится ниже ограничения, скважина переводится в режим bhp, после чего система решается повторно.")
    add_body_paragraph(doc, "В уравнение давления для скважин добавлен учет начальной карты давления. Поэтому текущая строка СЛАУ для i-й скважины записывается так:")
    add_formula(doc, "pavg - pwf,ᵢ + Σⱼ Mᵢⱼ qⱼ = -[pinit(xᵢ) - pinit,avg]")
    add_body_paragraph(doc, "Карта давления после решения системы восстанавливается как сумма среднего давления, начального смещения и вкладов всех 12 скважин. Таким образом, по сравнению с прошлым отчетом изменились только расчетная дискретизация и организация СЛАУ: суперпозиция масштабирована на 12 скважин, добавлена временная зависимость q_D, а переключение ограничений реализовано через controls вместо явной записи с F_i.")

    add_heading(doc, "3. Результаты вычислений и сравнение с tNavigator", level=1)
    add_body_paragraph(doc, "Расчет выполнен для доступных временных точек tNavigator. В основной текст вынесены 1, 5, 10, 120 и 288 месяцев, чтобы показать начальный участок, ранний переходный период и долгосрочную динамику без перегрузки отчета однотипными картами.")
    add_body_paragraph(doc, "Среднее начальное давление по поровому объему составляет 217.845 бар. Суммарный целевой поверхностный дебит для 12 скважин равен 120000 м3/сут, суммарный пластовый дебит - 180000 м3/сут. В расчете использована откалиброванная сжимаемость c_t = 0.030286 1/бар.")
    add_comparison_table(doc)
    add_picture(doc, error_chart, 15.5, "Рисунок 2 - Динамика средней абсолютной ошибки и RMSE по всем доступным месяцам")
    add_body_paragraph(doc, "По таблице и графику видно, что на ранних сроках средняя ошибка остается малой: для 1 месяца средняя абсолютная ошибка составляет 0.081 бар, для 5 месяцев - 0.141 бар, для 10 месяцев - 0.226 бар. К 120 месяцам ошибка возрастает до 2.561 бар, а к 288 месяцам - до 7.090 бар.")
    add_body_paragraph(doc, "Ниже приведены карты давления PSS-DTOF, tNavigator и абсолютной ошибки для выбранных моментов времени.")
    figure_number = 3
    for month in SELECTED_MONTHS:
        add_picture(doc, comparison_paths[month], 16.5, f"Рисунок {figure_number} - Сравнение PSS-DTOF, tNavigator и абсолютной ошибки через {month} мес.")
        figure_number += 1
    add_body_paragraph(doc, "На ранних сроках PSS-DTOF достаточно хорошо воспроизводит средний уровень давления, но локальная максимальная ошибка остается заметной вблизи отдельных скважин. На поздних сроках различие по среднему давлению растет, что указывает на необходимость дальнейшей калибровки параметров и проверки предположений суперпозиционной модели.")

    add_heading(doc, "4. Выводы", level=1)
    add_body_paragraph(doc, "В ходе работы реализована постановка PSS-DTOF суперпозиции для группы из 12 скважин на модели карбонатного коллектора Costa.")
    add_bullets(doc, [
        "построена матрица взаимного влияния M_{i,j}(t), учитывающая вклад каждой скважины-источника в точке каждой скважины-наблюдателя;",
        "в расчет включены формулы для переходного и начального режимов через q_D(τ,t), A_k(t) и J^{-1}(t);",
        "получены карты давления и карты абсолютной ошибки относительно tNavigator;",
        "лучшее совпадение наблюдается на ранних сроках, где средняя ошибка меньше 0.3 бар до 10 месяцев;",
        "на поздних сроках ошибка увеличивается: около 2.56 бар по средней абсолютной ошибке на 120 месяце и около 7.09 бар на 288 месяце.",
    ])
    add_body_paragraph(doc, "Полученные результаты показывают, что суперпозиционный PSS-DTOF подход применим для быстрого построения карт давления и анализа общей динамики, но для долгосрочного прогноза требуется дополнительная настройка и проверка причин накопления ошибки.")

    add_heading(doc, "Список используемых источников", level=1)
    sources = [
        "Costa Gomes J., Maschio M., Schiozer D. The design of an open-source carbonate reservoir model (the COSTA model) // Petroleum Geoscience. - 2022. - Vol. 28, No. 4.",
        "Матвеев И. В., Попов Д. А., Новгородова А. Н., Рукавишников В. С., Шишаев Г. Ю. Практическая реализация метода PseudoSteadyState (PSS) для моделирования пространственного распределения давления.",
        "King M. J., Wang Z., Datta-Gupta A. Asymptotic Solutions of the Diffusivity Equation and Their Applications. SPE 180149-MS, 2016.",
        "Nakajima K., King M. Development and Application of Fast Simulation Based on the PSS Pressure as a Spatial Coordinate. SPE-206085-MS.",
        "Nelson R. A. Geologic Analysis of Naturally Fractured Reservoirs. Gulf Professional Publishing, 2001.",
    ]
    for idx, source in enumerate(sources, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        r = p.add_run(f"{idx}. ")
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(14)
        r = p.add_run(source)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(14)

    doc.save(str(OUT_DOCX))


def main() -> None:
    ASSETS.mkdir(exist_ok=True)
    comparison_paths = extract_notebook_figures()
    model_maps = make_model_maps()
    error_chart = draw_error_chart()
    build_doc(model_maps, comparison_paths, error_chart)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
